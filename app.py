"""
Streamlit UI for Insurance Claim Assistant.

A multi-agent system that analyzes insurance claim denials and generates appeal letters.
"""

import os
from dotenv import load_dotenv

# Load environment variables FIRST
load_dotenv()

import streamlit as st
from pathlib import Path
import io
import urllib.parse

import config
from tools.pdf_parser import extract_text_from_pdf, extract_claim_info
from tools.denial_codes import analyze_denial_codes, format_denial_analysis_report
from crew import run_claim_analysis

# Word document generation
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def separate_letter_and_next_steps(full_result: str) -> tuple:
    """
    Separate the appeal letter content from the 'Next Steps' guidance.
    Returns (letter_content, next_steps_content)
    """
    # Common markers where next steps begin
    next_steps_markers = [
        "**Next Steps:**",
        "**Next Steps**",
        "## Next Steps",
        "### Next Steps",
        "**Final Notes:**",
        "**Final Notes**",
        "## Final Notes",
    ]
    
    letter_content = full_result
    next_steps_content = ""
    
    for marker in next_steps_markers:
        if marker in full_result:
            parts = full_result.split(marker, 1)
            letter_content = parts[0].strip()
            next_steps_content = marker + parts[1]
            break
    
    return letter_content, next_steps_content


def generate_word_doc(letter_content: str, patient_name: str = "Policyholder") -> bytes:
    """Generate a Word document from the appeal letter content."""
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # Title
    title = doc.add_heading("Appeal Letter - Insurance Claim", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content - handle markdown formatting
    lines = letter_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('**') and line.endswith('**'):
            # Bold text (heading-like)
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
        elif line.startswith('###'):
            doc.add_heading(line.replace('#', '').strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line.replace('#', '').strip(), level=2)
        elif line.startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line[0:3].rstrip('.').isdigit() and '. ' in line[:4]:
            # Numbered list
            doc.add_paragraph(line.split('. ', 1)[1] if '. ' in line else line, style='List Number')
        else:
            # Regular paragraph - clean up any remaining markdown
            clean_line = line.replace('**', '').replace('*', '')
            doc.add_paragraph(clean_line)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


def create_mailto_link(email: str, subject: str, body: str) -> str:
    """Create a mailto link for email clients."""
    params = {
        "subject": subject,
        "body": body[:2000]  # Limit body length for URL compatibility
    }
    query = urllib.parse.urlencode(params, quote_via=urllib.parse.quote)
    return f"mailto:{email}?{query}"


# Page configuration
st.set_page_config(
    page_title="Insurance Claim Assistant",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)


def render_header():
    """Render the main header."""
    st.markdown("""
    <div class="main-header">
        <h1>📋 Insurance Claim Assistant (India)</h1>
        <p>AI-powered analysis of health insurance claim rejections and appeal letter generation</p>
        <p style="font-size: 0.9rem; opacity: 0.8;">IRDAI Guidelines | Insurance Ombudsman | Consumer Forum</p>
    </div>
    """, unsafe_allow_html=True)


def render_sidebar():
    """Render the sidebar with information and settings."""
    with st.sidebar:
        # Home button - always visible
        if st.button("🏠 Go to Home", use_container_width=True, type="secondary"):
            for key in ["document_text", "claim_info", "denial_codes", 
                       "analysis_result", "patient_info", "policy_text",
                       "auto_populated", "show_auto_fill_toast"]:
                if key in st.session_state:
                    del st.session_state[key]
            st.session_state.step = 1
            st.rerun()
        
        st.divider()
        
        st.header("📖 How It Works")
        
        st.markdown("""
        **1. Upload Documents**
        Upload your claim rejection letter (PDF) or paste text
        
        **2. Add Details**
        Provide your information and policy documents
        
        **3. AI Analysis**
        Our 6 AI agents will:
        - Extract claim & rejection details
        - Analyze policy as per IRDAI guidelines
        - Review rejection validity
        - Develop appeal strategy
        - Write appeal letter (Hindi/English)
        - Quality review
        
        **4. Get Results**
        Receive appeal letter for GRO, Ombudsman, or Consumer Forum
        """)
        
        st.divider()
        
        st.header("⚙️ Settings")
        
        # Check API key
        if config.MISTRAL_API_KEY:
            st.success("✅ Mistral API Key configured")
        else:
            st.error("❌ Mistral API Key missing")
            st.info("Add MISTRAL_API_KEY to your .env file")
        
        # LangSmith status
        if config.LANGSMITH_ENABLED:
            st.success("✅ LangSmith tracing enabled")
        else:
            st.info("💡 Enable LangSmith for debugging")
        
        st.divider()
        
        st.header("📚 Resources (India)")
        st.markdown("""
        - [IRDAI IGMS Portal](https://igms.irda.gov.in)
        - [Insurance Ombudsman](https://www.cioins.co.in)
        - [IRDAI Guidelines](https://www.irdai.gov.in)
        - [Consumer Helpline: 155255](tel:155255)
        """)


def initialize_session_state():
    """Initialize session state variables."""
    if "step" not in st.session_state:
        st.session_state.step = 1
    if "document_text" not in st.session_state:
        st.session_state.document_text = ""
    if "claim_info" not in st.session_state:
        st.session_state.claim_info = {}
    if "denial_codes" not in st.session_state:
        st.session_state.denial_codes = []
    if "analysis_result" not in st.session_state:
        st.session_state.analysis_result = None
    if "patient_info" not in st.session_state:
        st.session_state.patient_info = {}
    if "auto_populated" not in st.session_state:
        st.session_state.auto_populated = False
    if "show_auto_fill_toast" not in st.session_state:
        st.session_state.show_auto_fill_toast = False


def render_step_1():
    """Step 1: Upload documents."""
    st.header("Step 1: Upload Your Documents")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📄 Denial Letter (Required)")
        denial_file = st.file_uploader(
            "Upload your denial letter or EOB",
            type=["pdf"],
            key="denial_upload",
            help="Upload the denial letter or Explanation of Benefits (EOB) you received"
        )
        
        if denial_file:
            with st.spinner("📄 Extracting text from document..."):
                file_bytes = io.BytesIO(denial_file.read())
                text = extract_text_from_pdf(file_bytes)
                
                if text:
                    st.session_state.document_text = text
                    st.success(f"✅ Extracted {len(text)} characters from document")
                    
                    with st.spinner("🤖 AI is analyzing and extracting claim details..."):
                        st.session_state.claim_info = extract_claim_info(text)
                    
                    # Extract denial codes
                    if st.session_state.claim_info.get("denial_codes"):
                        st.session_state.denial_codes = st.session_state.claim_info["denial_codes"]
                    
                    # Mark as auto-populated
                    st.session_state.auto_populated = True
                    st.session_state.show_auto_fill_toast = True
                    
                    # Show what was extracted
                    filled = [k for k, v in st.session_state.claim_info.items() if v and k != "denial_codes"]
                    if filled:
                        st.success(f"🎉 AI extracted: {', '.join(filled)}")
                    
                    with st.expander("Preview extracted text"):
                        st.text(text[:2000] + "..." if len(text) > 2000 else text)
                else:
                    st.error("Could not extract text from PDF. Please try a different file.")
    
    with col2:
        st.subheader("📑 Insurance Policy (Optional)")
        policy_file = st.file_uploader(
            "Upload your insurance policy",
            type=["pdf"],
            key="policy_upload",
            help="Upload your policy document for more accurate coverage analysis"
        )
        
        if policy_file:
            with st.spinner("Extracting policy text..."):
                file_bytes = io.BytesIO(policy_file.read())
                policy_text = extract_text_from_pdf(file_bytes)
                
                if policy_text:
                    st.session_state.policy_text = policy_text
                    st.success(f"✅ Extracted policy document")
                else:
                    st.warning("Could not extract text from policy PDF")
    
    # Manual text input option
    st.divider()
    st.subheader("📝 Or Paste Text Directly")
    
    manual_text = st.text_area(
        "Paste your denial letter text here",
        height=200,
        placeholder="Copy and paste the text from your denial letter...",
        help="If you can't upload a PDF, paste the text directly"
    )
    
    if manual_text and manual_text != st.session_state.get("last_manual_text", ""):
        st.session_state.last_manual_text = manual_text
        st.session_state.document_text = manual_text
        
        with st.spinner("🤖 AI is analyzing and extracting claim details..."):
            st.session_state.claim_info = extract_claim_info(manual_text)
        
        if st.session_state.claim_info.get("denial_codes"):
            st.session_state.denial_codes = st.session_state.claim_info["denial_codes"]
        
        # Mark as auto-populated
        st.session_state.auto_populated = True
        st.session_state.show_auto_fill_toast = True
        
        # Show what was extracted
        filled = [k for k, v in st.session_state.claim_info.items() if v and k != "denial_codes"]
        if filled:
            st.success(f"🎉 AI extracted: {', '.join(filled)}")
    
    # Continue button
    if st.session_state.document_text:
        if st.button("Continue to Step 2 →", type="primary"):
            st.session_state.step = 2
            st.rerun()


def render_step_2():
    """Step 2: Review extracted info and add patient details."""
    st.header("Step 2: Review & Add Details")
    
    # Show auto-fill toast notification
    if st.session_state.get("show_auto_fill_toast", False):
        # Count auto-filled fields
        claim_info = st.session_state.claim_info
        filled_fields = [k for k, v in claim_info.items() if v and k != "denial_codes"]
        if st.session_state.denial_codes:
            filled_fields.append("denial_codes")
        
        if filled_fields:
            st.toast(f"✅ Auto-filled {len(filled_fields)} fields from your document! Please review and complete the remaining fields.", icon="🎉")
        st.session_state.show_auto_fill_toast = False
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📋 Claim Information")
        st.caption("ℹ️ Fields marked with ✓ were auto-extracted from your document")
        
        claim_info = st.session_state.claim_info
        
        # Helper to show if field was auto-filled
        def field_label(label, field_name):
            if claim_info.get(field_name):
                return f"{label} ✓"
            return label
        
        # Editable claim fields
        claim_number = st.text_input(
            field_label("Claim Number", "claim_number"),
            value=claim_info.get("claim_number", ""),
            help="The claim number from your denial letter"
        )
        
        policy_number = st.text_input(
            field_label("Policy Number", "policy_number"),
            value=claim_info.get("policy_number") or claim_info.get("member_id", ""),
        )
        
        insurer_name = st.text_input(
            field_label("Insurance Company", "insurer_name"),
            value=claim_info.get("insurer_name", ""),
            help="e.g., Star Health, ICICI Lombard, HDFC ERGO"
        )
        
        hospital_name = st.text_input(
            field_label("Hospital Name", "hospital_name"),
            value=claim_info.get("hospital_name", ""),
        )
        
        tpa_name = st.text_input(
            field_label("TPA Name", "tpa_name"),
            value=claim_info.get("tpa_name", ""),
            help="Third Party Administrator (if applicable)"
        )
        
        col1a, col1b = st.columns(2)
        with col1a:
            admission_date = st.text_input(
                field_label("Admission Date", "admission_date"),
                value=claim_info.get("admission_date") or claim_info.get("service_date", ""),
            )
        with col1b:
            discharge_date = st.text_input(
                field_label("Discharge Date", "discharge_date"),
                value=claim_info.get("discharge_date", ""),
            )
        
        claim_amount = st.text_input(
            field_label("Claim Amount (₹)", "claim_amount"),
            value=claim_info.get("claim_amount") or claim_info.get("billed_amount", ""),
        )
        
        denial_reason = st.text_area(
            field_label("Denial Reason", "denial_reason"),
            value=claim_info.get("denial_reason", ""),
            height=100,
        )
        
        # Denial codes
        st.subheader("🔢 Denial Codes")
        
        denial_codes_str = st.text_input(
            "Denial Codes (comma-separated)" + (" ✓" if st.session_state.denial_codes else ""),
            value=", ".join(st.session_state.denial_codes),
            help="e.g., PED-001, WP-002, PA-001"
        )
        
        if denial_codes_str:
            codes = [c.strip() for c in denial_codes_str.split(",") if c.strip()]
            st.session_state.denial_codes = codes
            
            # Show denial code analysis
            if codes:
                analysis = analyze_denial_codes(codes)
                
                st.markdown(f"""
                <div class="info-card">
                    <strong>Quick Analysis:</strong><br>
                    Appeal Likelihood: <strong>{analysis['overall_appeal_likelihood']}</strong><br>
                    {analysis['summary']}
                </div>
                """, unsafe_allow_html=True)
    
    with col2:
        st.subheader("👤 Your Information")
        st.caption("ℹ️ Please fill in your details for the appeal letter")
        
        # Check if patient name was auto-extracted
        patient_name = st.text_input(
            field_label("Your Full Name", "patient_name"),
            value=claim_info.get("patient_name") or st.session_state.patient_info.get("name", "")
        )
        patient_address = st.text_area("Your Address", height=100, value=st.session_state.patient_info.get("address", ""))
        patient_phone = st.text_input("Phone Number", value=st.session_state.patient_info.get("phone", ""))
        patient_email = st.text_input("Email Address", value=st.session_state.patient_info.get("email", ""))
        
        st.session_state.patient_info = {
            "name": patient_name,
            "address": patient_address,
            "phone": patient_phone,
            "email": patient_email,
        }
        
        # Update claim info with all the extracted and edited values
        st.session_state.claim_info.update({
            "claim_number": claim_number,
            "policy_number": policy_number,
            "insurer_name": insurer_name,
            "hospital_name": hospital_name,
            "tpa_name": tpa_name,
            "admission_date": admission_date,
            "discharge_date": discharge_date,
            "claim_amount": claim_amount,
            "denial_reason": denial_reason,
        })
    
    # Navigation
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("← Back to Step 1"):
            st.session_state.step = 1
            st.rerun()
    
    with col2:
        if st.button("Analyze & Generate Appeal →", type="primary"):
            if not st.session_state.document_text:
                st.error("Please upload a document first")
            else:
                st.session_state.step = 3
                st.rerun()


def render_step_3():
    """Step 3: Run analysis and show results."""
    st.header("Step 3: AI Analysis")
    
    if st.session_state.analysis_result is None:
        st.info("🤖 Our AI agents are analyzing your claim and generating an appeal letter...")
        
        # Progress display
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        agents = [
            "Document Analyzer - Extracting claim details...",
            "Policy Expert - Analyzing coverage...",
            "Denial Reviewer - Checking validity...",
            "Appeal Strategist - Building strategy...",
            "Letter Writer - Drafting appeal...",
            "Quality Reviewer - Final polish...",
        ]
        
        try:
            # Run the crew
            with st.spinner("Running AI analysis..."):
                for i, agent in enumerate(agents):
                    status_text.text(f"🔄 {agent}")
                    progress_bar.progress((i + 1) / len(agents))
                
                result = run_claim_analysis(
                    document_text=st.session_state.document_text,
                    denial_codes=st.session_state.denial_codes,
                    policy_text=st.session_state.get("policy_text"),
                    patient_info=st.session_state.patient_info,
                )
                
                st.session_state.analysis_result = str(result)
                progress_bar.progress(100)
                status_text.text("✅ Analysis complete!")
                st.rerun()
                
        except Exception as e:
            st.error(f"Error during analysis: {str(e)}")
            st.info("Please check your API key and try again.")
            
            if st.button("← Back to Step 2"):
                st.session_state.step = 2
                st.rerun()
    
    else:
        # Show results
        st.success("✅ Analysis Complete!")
        
        # Display the appeal letter
        st.subheader("📄 Your Appeal Letter")
        
        result = st.session_state.analysis_result
        patient_name = st.session_state.patient_info.get("name", "Policyholder")
        patient_email = st.session_state.patient_info.get("email", "")
        claim_number = st.session_state.claim_info.get("claim_number", "")
        insurer_name = st.session_state.claim_info.get("insurer_name", "Insurance Company")
        
        # Separate letter content from Next Steps guidance
        letter_content, ai_next_steps = separate_letter_and_next_steps(result)
        
        # Display in a nice format
        st.markdown("""
        <div class="success-card">
            <strong>Your appeal letter has been generated!</strong>
            Review it below, make any needed edits, then download or send via email.
        </div>
        """, unsafe_allow_html=True)
        
        # Editable text area with ONLY the letter content (no next steps)
        edited_letter = st.text_area(
            "Appeal Letter (editable)",
            value=letter_content,
            height=500,
        )
        
        # Download and action options
        st.subheader("📤 Export Options")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Word download
            if DOCX_AVAILABLE:
                try:
                    docx_bytes = generate_word_doc(edited_letter, patient_name)
                    if docx_bytes:
                        st.download_button(
                            label="📄 Download Word",
                            data=docx_bytes,
                            file_name=f"appeal_letter_{claim_number or 'claim'}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                except Exception as e:
                    st.download_button(
                        label="📄 Download TXT",
                        data=edited_letter,
                        file_name=f"appeal_letter_{claim_number or 'claim'}.txt",
                        mime="text/plain",
                    )
            else:
                st.download_button(
                    label="📄 Download TXT",
                    data=edited_letter,
                    file_name=f"appeal_letter_{claim_number or 'claim'}.txt",
                    mime="text/plain",
                )
        
        with col2:
            # Email button
            email_subject = f"Appeal against Claim Rejection - {claim_number}" if claim_number else "Appeal against Claim Rejection"
            email_body = f"Dear Sir/Madam,\n\nPlease find below my appeal letter.\n\n---\n\n{edited_letter[:1500]}...\n\n(Full letter attached)"
            
            # Create mailto link
            mailto_link = create_mailto_link(
                email="grievance@" + insurer_name.lower().replace(" ", "") + ".com" if insurer_name else "",
                subject=email_subject,
                body=email_body
            )
            
            st.markdown(f"""
            <a href="{mailto_link}" target="_blank" style="
                display: inline-block;
                padding: 0.5rem 1rem;
                background-color: #4CAF50;
                color: white;
                text-decoration: none;
                border-radius: 5px;
                font-size: 14px;
            ">📧 Send via Email</a>
            """, unsafe_allow_html=True)
        
        # AI-Generated Next Steps (from the analysis)
        st.divider()
        
        if ai_next_steps:
            with st.expander("🤖 AI-Generated Guidance & Next Steps", expanded=True):
                st.markdown(ai_next_steps)
        
        # Standard next steps checklist
        st.subheader("📌 Your Action Checklist")
        
        st.markdown("""
        1. **Review the letter** - Make sure all information is accurate
        2. **Fill in placeholders** - Replace any [BRACKETS] with actual information
        3. **Gather documents** - Collect all supporting documents mentioned in enclosures
        4. **Print and sign** - Sign the letter before sending
        5. **Send via certified mail** - Keep proof of delivery
        6. **Keep copies** - Save copies of everything you send
        7. **Note the deadline** - Mark your calendar for response deadline (15 days for GRO, 30 days for Ombudsman)
        """)
        
        # Quick escalation guide
        st.subheader("📞 Escalation Contacts")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            **Level 1: GRO**
            - Contact your insurer's Grievance Redressal Officer
            - Response time: 15 days
            """)
        
        with col2:
            st.markdown("""
            **Level 2: IGMS Portal**
            - [igms.irda.gov.in](https://igms.irda.gov.in)
            - Helpline: 155255
            """)
        
        with col3:
            st.markdown("""
            **Level 3: Ombudsman**
            - [cioins.co.in](https://www.cioins.co.in)
            - For claims up to ₹50 lakhs
            """)
        
        # Warning
        st.markdown("""
        <div class="warning-card">
            <strong>⚠️ Important Disclaimer</strong><br>
            This tool provides general guidance only and is not legal advice. 
            For complex cases, consider consulting with a patient advocate or attorney 
            specializing in insurance claims.
        </div>
        """, unsafe_allow_html=True)


def main():
    """Main application entry point."""
    initialize_session_state()
    
    # Apply static CSS
    st.markdown("""
    <style>
        /* Header styling */
        .main-header {
            background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 100%);
            padding: 2rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            color: white;
            text-align: center;
        }
        
        .main-header h1 {
            margin: 0;
            font-size: 2.5rem;
        }
        
        .main-header p {
            margin: 0.5rem 0 0 0;
            opacity: 0.9;
        }
        
        /* Card styling */
        .info-card {
            background: #f8f9fa;
            border: 1px solid #e9ecef;
            border-radius: 8px;
            padding: 1.5rem;
            margin-bottom: 1rem;
        }
        
        .success-card {
            background: #d4edda;
            border: 1px solid #c3e6cb;
            border-radius: 8px;
            padding: 1.5rem;
            margin-bottom: 1rem;
        }
        
        .warning-card {
            background: #fff3cd;
            border: 2px solid #ffc107;
            border-radius: 8px;
            padding: 1.5rem;
            margin-bottom: 1rem;
            color: #856404 !important;
        }
        
        .warning-card strong {
            color: #664d03 !important;
        }
        
        /* Hide Streamlit branding */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
    </style>
    """, unsafe_allow_html=True)
    
    render_header()
    render_sidebar()
    
    # Check for API key
    if not config.MISTRAL_API_KEY:
        st.error("⚠️ Mistral API Key not configured!")
        st.info("""
        To use this application:
        1. Get an API key from [Mistral AI](https://console.mistral.ai/)
        2. Create a `.env` file in the project directory
        3. Add: `MISTRAL_API_KEY=your-key-here`
        4. Restart the application
        """)
        return
    
    # Render current step
    if st.session_state.step == 1:
        render_step_1()
    elif st.session_state.step == 2:
        render_step_2()
    elif st.session_state.step == 3:
        render_step_3()


if __name__ == "__main__":
    main()
